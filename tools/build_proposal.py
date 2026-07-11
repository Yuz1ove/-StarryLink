from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OFFICIAL = ROOT / "official"
OUT = ROOT / "submission"

TEAM_NAME = "韌訊小隊"
WORK_TITLE = "星夜守望者 AI 韌性通訊編排平台"
DOCX_OUT = OUT / f"{TEAM_NAME}_作品提案規劃書.docx"


def set_font(run, size=12, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(p, after=6, line=1.2, align=None):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    for run in p.runs:
        if run.text:
            set_font(run)


def clear_cell(cell):
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)


def add_p(cell, text="", size=12, bold=False, color=None, after=6, line=1.2, align=None):
    p = cell.add_paragraph()
    style_paragraph(p, after=after, line=line, align=align)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(cell, text):
    p = add_p(cell, text, size=13, bold=True, color="1F4D78", after=4, line=1.15)
    p.paragraph_format.space_before = Pt(8)
    return p


def add_body(cell, text):
    return add_p(cell, text, size=12, after=6, line=1.25)


def add_bullets(cell, items):
    for item in items:
        p = cell.add_paragraph(style=None)
        style_paragraph(p, after=3, line=1.18)
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        r = p.add_run("• ")
        set_font(r, size=12)
        r = p.add_run(item)
        set_font(r, size=12)


def add_numbered(cell, items):
    for idx, item in enumerate(items, 1):
        p = cell.add_paragraph(style=None)
        style_paragraph(p, after=3, line=1.18)
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        r = p.add_run(f"{idx}. ")
        set_font(r, size=12)
        r = p.add_run(item)
        set_font(r, size=12)


def add_compact_kv(cell, rows):
    for key, value in rows:
        p = cell.add_paragraph(style=None)
        style_paragraph(p, after=3, line=1.12)
        p.paragraph_format.left_indent = Cm(0.2)
        r = p.add_run(f"{key}：")
        set_font(r, size=12, bold=True, color="0B2545")
        r = p.add_run(value)
        set_font(r, size=12)


def add_footer_page_number(doc):
    section = doc.sections[0]
    footer = section.footer
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("第 ")
    set_font(r, size=10)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_end)
    r = p.add_run(" 頁")
    set_font(r, size=10)


def cell_text(cell, text):
    clear_cell(cell)
    p = cell.add_paragraph()
    style_paragraph(p, after=0, line=1.15)
    r = p.add_run(text)
    set_font(r, size=12)


def build():
    OUT.mkdir(exist_ok=True)
    template = next(OFFICIAL.glob("02*.docx"))
    doc = Document(template)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    for p in doc.paragraphs:
        style_paragraph(p, after=3, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER if p.text.strip() in {"中華電信智慧創新應用大賽", "作品提案規劃書"} else None)
        for run in p.runs:
            if run.text.strip() == "中華電信智慧創新應用大賽":
                set_font(run, size=18, bold=True, color="0B2545")
            elif run.text.strip() == "作品提案規劃書":
                set_font(run, size=16, bold=True, color="1F4D78")
            else:
                set_font(run, size=12)

    tbl = doc.tables[0]
    cell_text(tbl.cell(0, 1), "□校園組  ☑社會組")
    cell_text(tbl.cell(1, 1), WORK_TITLE)
    cell_text(tbl.cell(2, 1), TEAM_NAME)
    cell_text(
        tbl.cell(3, 1),
        "□智慧醫療    □智慧環境\n□智慧交通    □智慧金融\n□智慧製造    □智慧生活\n□智慧文化    ☑智慧通信",
    )
    cell_text(tbl.cell(4, 1), "☑ 本作品使用海地星空應用，並申請「海地星空應用獎」")

    intro = (
        "星夜守望者是面向家庭、社區、企業據點與地方救災單位的 AI 韌性通訊助理。"
        "系統結合行動、固網寬頻、衛星、微波、海纜等「海地星空」多元通訊能力，以 AI Agent 判讀使用者情境、網路狀態與事件嚴重度，自動選擇低頻寬文字、語音摘要、群組回報或衛星備援等通訊策略。"
        "它把平時的生活助理、防災通知、家人平安回報、企業營運續航與中華電信網路服務整合為可落地的訂閱服務，降低災害與斷訊時的資訊落差，並提升中華電信在智慧通信與韌性城市市場的服務價值。"
    )
    cell_text(tbl.cell(5, 1), intro)

    explain = tbl.cell(6, 1)
    clear_cell(explain)

    add_heading(explain, "壹、主題說明")
    add_body(
        explain,
        "本作品選擇「智慧通信」主題，核心問題是一般用戶在災害、尖峰壅塞、基地台維修、偏鄉或移動場景中，常不知道哪一種通訊方式最可靠，也缺少能把家人、企業同仁、社區與救災資訊整合在一起的即時助理。現有緊急通報多偏向單向廣播；一般通訊 App 又仰賴既有網路條件，當網路品質下降時，資訊容易分散、重複或延遲。",
    )
    add_body(
        explain,
        "星夜守望者以中華電信多元網路資源為基底，建立「平時好用、急時能用、斷點可續」的通訊服務。平時提供生活與家庭安全助理；事件發生時，AI 依據使用者身分、地點、通訊可用性與事件類型，切換為精簡訊息、語音播報、群組平安回報、企業值班通報或衛星備援指引。目標對象包含家庭用戶、銀髮照護家庭、社區管委會、中小企業據點、戶外工作者與地方防救災合作單位。",
    )

    add_heading(explain, "貳、作品特色")
    add_bullets(
        explain,
        [
            "AI 韌性通訊決策：以規則引擎搭配 AI Agent，依事件嚴重度、網路品質、收件對象與可用頻寬，推薦或自動切換訊息格式與通訊路徑。",
            "海地星空整合：把行動網路、固網寬頻、衛星、微波、海纜等多元網路視為服務能力池，將中華電信的網路優勢轉化為使用者可感知的韌性體驗。",
            "低頻寬優先設計：災害或壅塞時先傳送文字、狀態碼、位置摘要與必要語音，再補送完整內容，提升弱網環境下的訊息抵達率。",
            "生活助理與緊急模式雙軌：平時可查詢家人平安、設備狀態、社區公告與行程提醒；急時自動進入事件面板，集中處理求助、回報、確認與指引。",
            "企業與社區可擴充：支援值班群組、據點清冊、員工安全回報、停電/斷網應變流程與地方單位公告串接，具 B2C、B2B2C、B2G 商轉空間。",
        ],
    )

    add_heading(explain, "參、設計理念")
    add_body(
        explain,
        "設計理念是「用 AI 把通訊服務從連線工具升級為韌性決策服務」。使用者真正需要的不是知道所有網路技術細節，而是在關鍵時刻確定訊息能被送達、被理解、被回覆。作品以三個原則設計：第一，先保命再補充，所有緊急模式都以最小必要資訊為優先；第二，先人後系統，介面避免複雜設定，讓家庭長者或第一線人員也能快速回報；第三，先落地再擴大，初期以既有手機 App、簡訊、語音、Web 管理台與雲端服務實作，不需等待新硬體普及。",
    )
    add_body(
        explain,
        "與一般聊天機器人不同，本作品不只是回答問題，而是依據網路狀態與任務緊急度做通訊編排；與一般防災 App 不同，本作品不只推播公告，而是把個人、家庭、企業與社區的回報閉環建立起來；與單一衛星或備援方案不同，本作品把多網路能力抽象為可營運的服務層，可隨中華電信既有產品組合逐步導入。",
    )

    add_heading(explain, "肆、架構說明")
    add_body(explain, "系統由使用者端、AI 編排層、通訊服務層、資料治理層與營運管理台五部分組成；MVP 以可展示的 Web 控制台、網路狀態模擬器與通道決策引擎呈現。")
    add_numbered(
        explain,
        [
            "使用者端：手機 App、Web 入口、簡訊互動與語音互動，提供一鍵報平安、緊急求助、家庭/團隊狀態面板與管理台。",
            "事件 API：POST /events 建立事件，欄位包含 event_type、severity、location_risk、group_size、ack_target、privacy_level、battery_level、network_snapshot。",
            "AI 編排層：先做事件分級與訊息壓縮，再依通道評分選擇傳送路徑；RAG 知識庫可放防災 SOP、企業應變手冊、地方公告與服務 FAQ。",
            "通訊服務層：抽象化 Push、SMS、語音 IVR、衛星簡訊、企業專線、微波/固定備援等通道，統一輸出 route_plan、payload_format、fallback_order。",
            "資料治理層：保存最小必要個資、事件紀錄、同意設定、角色權限與稽核紀錄；緊急事件資料預設留存 30 天，可依企業或公部門專案調整。",
            "營運管理台：即時顯示事件等級、通道分數、送達率、回覆率、未確認名單、SLA 倒數與備援通道啟用紀錄。",
        ],
    )
    add_compact_kv(
        explain,
        [
            ("通道評分公式", "score = 0.30*availability + 0.20*signal + 0.15*latency_score + 0.15*loss_score + 0.10*cost_score + 0.10*policy_match。分數低於 60 分不作為主通道；低於 45 分只列為備援。"),
            ("事件分級", "severity 1-2 為一般提醒，3 為注意事件，4 為緊急確認，5 為生命安全或營運中斷；severity >= 4 時自動啟用至少兩條通道。"),
            ("低頻寬閾值", "available_bandwidth < 64 kbps 時改為 160 字以內文字；< 16 kbps 時只保留狀態碼、座標、時間與回覆選項；完全斷網時排入衛星或語音備援。"),
            ("SLA 目標", "家庭平安回報 5 分鐘內取得 80% 回覆；企業據點 10 分鐘內取得 90% 值班確認；社區公告 30 分鐘內產出未確認名單。"),
        ],
    )
    add_body(
        explain,
        "Demo 實作會提供三個可切換情境：颱風夜家庭平安回報、企業據點斷網應變、山區巡檢弱訊號。評審可調整頻寬、延遲、封包遺失、電量、事件等級與群組人數，畫面會即時計算主通道、備援順序、訊息壓縮後 payload、預估送達率與未確認名單。",
    )

    add_heading(explain, "伍、使用情境")
    add_bullets(
        explain,
        [
            "家庭安全：颱風夜網路壅塞時，使用者按下「一鍵報平安」，AI 將位置與狀態壓縮成低頻寬訊息，優先通知家人；若家人未回覆，系統提醒改用語音或簡訊。",
            "銀髮照護：長者不需操作複雜 App，可透過語音說「我沒事」或「需要協助」，系統自動轉成文字摘要並推送給照護者。",
            "企業營運續航：中小企業設定分店與值班群組，斷網或停電時由系統發出員工安全確認、客戶服務公告與備援聯繫流程。",
            "社區防災：管委會或里辦公室可發送分眾公告，系統統計回覆率與未確認名單，協助掌握資訊傳達缺口。",
            "偏鄉與戶外作業：在山區、海邊或移動巡檢場域，系統提示可用通訊方式與最可靠訊息格式，降低工作安全風險。",
        ],
    )

    add_heading(explain, "陸、商業模式")
    add_body(
        explain,
        "本作品可包裝為中華電信智慧通信加值服務，採多層訂閱與專案導入模式。B2C 端提供家庭守護月租，搭配行動門號、寬頻、智慧家庭或銀髮照護方案；B2B 端提供企業據點韌性通訊管理，依據點數、成員數與通訊量收費；B2G/社區端則可與地方防災、社區安全、偏鄉服務與企業 ESG 專案合作。對中華電信而言，作品可提升既有通訊資產的差異化價值，帶動加值服務 ARPU，並強化衛星、固網、行動與雲端服務的整體解決方案銷售。",
    )
    add_bullets(
        explain,
        [
            "家庭版：每月訂閱，包含家庭成員、安全清單、平安回報與緊急模式。",
            "社區版：以戶數或群組數計價，提供公告、回覆統計、弱網策略與管理台。",
            "企業版：以據點與人數計價，提供營運續航 SOP、值班群組、稽核報表與 SLA 指標。",
            "合作專案：與地方政府、防災單位、長照機構、保險或企業 ESG 專案共同推廣。",
        ],
    )

    add_heading(explain, "柒、預期成果")
    add_bullets(
        explain,
        [
            "初賽至決賽期間完成可展示 MVP：手機互動頁、Web 管理台、AI 編排 API、網路狀態模擬器與三個使用情境 Demo。",
            "驗證弱網情境下的訊息壓縮、通道切換與回報閉環，提出訊息抵達率、平均回覆時間、未確認名單比例等評估指標。",
            "完成商業化假設：家庭版、社區版、企業版三種定價草案與導入流程。",
            "形成可與中華電信既有行動、寬頻、雲端、衛星與企業服務接軌的合作藍圖。",
            "長期目標是在智慧城市、防災、長照與企業韌性市場建立可複製的智慧通信服務。",
        ],
    )

    add_heading(explain, "捌、開發工具及其他相關說明")
    add_bullets(
        explain,
        [
            "前端 Demo：HTML/CSS/JavaScript 或 React，呈現情境切換、參數滑桿、通道評分、事件時間線與訊息 payload 預覽。",
            "後端原型：Python FastAPI，核心端點包含 POST /route/plan、POST /notify/send、POST /ack、GET /events/{id}/dashboard。",
            "AI：LLM Agent 負責訊息摘要、收件者語氣轉換與低頻寬改寫；規則引擎負責可驗證的通道評分與閾值判斷，避免黑箱決策。",
            "資料模型：events、recipients、channels、route_plans、delivery_logs、ack_logs、consent_records、audit_logs 八張主表即可支撐 MVP。",
            "通訊整合：Demo 階段以模擬 API 呈現行動、固網、簡訊、語音、推播與衛星備援；正式合作階段再串接中華電信可開放之服務介面。",
            "資安與隱私：採最小化蒐集、角色權限、事件留存期限、可撤回同意與操作稽核。緊急模式僅使用必要資料，避免將敏感位置長期保存。",
        ],
    )

    add_heading(explain, "拾、Demo 展示規格")
    add_body(
        explain,
        "本作品可展示的最小功能為「事件建立 -> 通道評分 -> 訊息壓縮 -> 多通道送出 -> 回覆追蹤」。展示不需真實串接電信核心網，先以模擬器呈現決策邏輯與營運畫面，待進入合作驗證階段再替換為正式 API。",
    )
    add_compact_kv(
        explain,
        [
            ("輸入參數", "severity、bandwidth_kbps、latency_ms、packet_loss、battery、recipient_count、ack_target、channel_availability。"),
            ("輸出結果", "primary_channel、fallback_order、compressed_payload、estimated_reach_rate、ack_deadline、operator_actions。"),
            ("可展示畫面", "情境面板、通道分數排行、訊息預覽、事件時間線、回覆統計、未確認名單、備援啟用原因。"),
            ("決賽延伸", "加入語音轉文字、實際簡訊測試帳號、地圖熱區、企業值班名冊匯入與管理者報表匯出。"),
        ],
    )

    add_heading(explain, "拾壹、參考資料")
    add_bullets(
        explain,
        [
            "2026 中華電信智慧創新應用大賽競賽辦法與作品提案規劃書模板。",
            "中華電信智慧創新應用大賽官方網站：https://cht5g.com.tw/",
            "獎金獵人活動頁：https://bhuntr.com/tw/competitions/ah6ooacwn1nx2ylujr",
        ],
    )

    add_heading(explain, "拾貳、補充說明")
    add_body(
        explain,
        "本提案不在文件中揭露參賽者姓名或參賽單位，符合官方模板提醒。正式報名時，參賽者仍需在報名系統與參賽同意書中填寫真實資料，並上傳身分證明文件。",
    )

    add_footer_page_number(doc)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    style_paragraph(p, after=p.paragraph_format.space_after.pt if p.paragraph_format.space_after else 3, line=p.paragraph_format.line_spacing or 1.15)
                    for run in p.runs:
                        if run.text:
                            size = 12
                            if run.bold and run.font.size and run.font.size.pt >= 13:
                                size = run.font.size.pt
                            set_font(run, size=size, bold=bool(run.bold), color=run.font.color.rgb.__str__() if run.font.color and run.font.color.rgb else None)

    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.title = WORK_TITLE
    doc.core_properties.subject = "2026 中華電信智慧創新應用大賽作品提案規劃書"
    doc.core_properties.comments = "Generated proposal draft without personal identifying information."

    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    build()
